import pdfplumber
import io
import docx
import re

from .skills_extractor import extract_skills
from .experience_extractor import extract_experience
from .project_extractor import extract_projects

def extract_text_from_file(file_bytes, mime_type):
    """
    Extract text from different file types with improved handling
    """
    if mime_type == 'application/pdf':
        return extract_text_from_pdf(file_bytes)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # DOCX handling
        return extract_text_from_docx(file_bytes)
    elif mime_type == 'application/msword':
        # Older .doc files - fallback to text extraction
        return extract_text_from_doc(file_bytes)
    elif mime_type == 'text/plain':
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('latin-1', errors='ignore')
    else:
        # Unexpected MIME types: try best-effort text decode
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

def extract_text_from_pdf(file_bytes):
    """
    Enhanced PDF text extraction with better error handling
    """
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up common PDF extraction artifacts
                        page_text = re.sub(r'\n+', '\n', page_text)  # Multiple newlines
                        page_text = re.sub(r' +', ' ', page_text)    # Multiple spaces
                        text += page_text + "\n"
                except Exception:
                    # Skip problematic pages
                    continue
            return text.strip()
    except Exception:
        # PDF may be corrupt, password-protected, or image-based
        return ""

def extract_text_from_docx(file_bytes):
    """
    Enhanced DOCX text extraction
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text += para_text + "\n"

        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += cell_text + " "

        return text.strip()
    except Exception:
        return ""

def extract_text_from_doc(file_bytes):
    """
    Fallback for older .doc files - basic text extraction
    """
    try:
        # For .doc files, we can't easily parse them with python-docx
        # Try to decode as text or return empty
        return file_bytes.decode('utf-8', errors='ignore')
    except Exception:
        return ""

def preprocess_text(text: str) -> str:
    """
    Preprocess extracted text to improve parsing accuracy
    """
    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r' +', ' ', text)

    # Remove common resume artifacts
    text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\d{1,2}/\d{1,2}/\d{4}', '', text)  # Remove dates
    text = re.sub(r'\d{4}-\d{2}-\d{2}', '', text)      # Remove ISO dates

    # Normalize bullet points
    text = re.sub(r'[•●○▪]', '•', text)

    return text.strip()

def parse_resume(file_bytes, mime_type='application/pdf'):
    """
    Parse resume and extract skills, experience, and projects with enhanced preprocessing
    """
    raw_text = extract_text_from_file(file_bytes, mime_type)
    processed_text = preprocess_text(raw_text)

    skills = extract_skills(processed_text)
    experience = extract_experience(processed_text)
    projects = extract_projects(processed_text)

    return {
        "raw_text": raw_text,
        "processed_text": processed_text,
        "skills": skills,
        "experience": experience,
        "projects": projects
    }
